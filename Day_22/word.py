"""wang

创建一个docx
"""
from docx import Document
from  docx.shared import Cm,Pt
from docx.document import Document as Doc

document =Document()
document.add_heading("快了个傻学python",0)
p = document.add_paragraph("python 是一个傻子学习的语言,它")
run = p.add_run("煞笔")
run.bold = True
run.font.size = Pt(18)
p.add_run("而且")
run = p.add_run("无敌煞笔")
run.font.size=Pt(12)
run.underline = True
p.add_run(".")

document.add_heading("Heading,level1",level=1)
document.add_paragraph("Intense quto",style="Intense Quote")
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'second item in ordered list', style='List Bullet'
)
# 添加有序列表
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
document.add_paragraph(
    'second item in ordered list', style='List Number'
)

# 添加图片（注意路径和图片必须要存在）
document.add_picture('resources/xuge.png', width=Cm(5.2))
document.add_section()
records =( ("煞笔","男","2006-6-4"),
           ("雷杰","男","2005-8-11")
           )
table = document.add_table(rows=1, cols=3)
table.style = 'Dark List'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '姓名'
hdr_cells[1].text = '性别'
hdr_cells[2].text = '出生日期'
for name ,sex,birthday in records:
    row_cells =table.add_row().cells
    row_cells[0].text=name
    row_cells[1].text=sex
    row_cells[2].text=birthday
document.add_page_break()
document.save("demo.docx")
